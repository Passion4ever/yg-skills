#!/usr/bin/env python3
"""把交底书 Markdown 渲染成 .docx。

版式沿用国内交底书的常见观感:正文宋体 11 pt、标题黑体、1.5 倍行距、首行缩进两字。
交底书无法定格式要求,下列常量按需改即可。

公式:`$$...$$` 独占一行的 LaTeX 渲染为**原生 Word 公式域**(OMML),代理师可双击编辑;
正文中的 `_{...}` / `^{...}` 渲染为上下标文本 run,用于行文里顺带提到的符号。

用法:
    python render_disclosure.py 交底书.md -o 交底书.docx
    python render_disclosure.py 交底书.md -o out.docx --landscape fig1 fig2
"""
from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor

# 校验器互相 import 会在技能目录里留下 __pycache__。短命 CLI 不需要字节码缓存，
# 关掉它，技能目录保持干净(打包时也就不会混进 .pyc)。
sys.dont_write_bytecode = True
sys.path.insert(0, str(Path(__file__).parent))
from math_to_omml import latex_to_omml  # noqa: E402

# ---- 版式常量 -------------------------------------------------------------
PAGE_W, PAGE_H = Mm(210), Mm(297)          # A4
MARGIN = Mm(25.4)
BODY_FONT_LATIN, BODY_FONT_CJK = 'Times New Roman', '宋体'
HEAD_FONT_LATIN, HEAD_FONT_CJK = 'Arial', '黑体'
BODY_PT = Pt(11)
HEAD_PT = {1: Pt(16), 2: Pt(14), 3: Pt(12), 4: Pt(12)}
LINE_SPACING = 1.5
FIRST_INDENT = Pt(22)                      # 约两个汉字
TABLE_PT = Pt(10)
CAPTION_PT = Pt(10.5)
HEADER_FILL = 'E8E8E8'
# 竖版页可用宽度 ≈ 159 mm;横版 ≈ 246 mm
FIT_MM = {'portrait': (155, 200), 'landscape': (240, 140)}

MD_IMG = re.compile(r'^!\[([^\]]*)\]\(([^)]+)\)$')
MD_HEAD = re.compile(r'^(#{1,4})\s+(.*)$')
MD_CAPTION = re.compile(r'^图\s*\d+\s{2}')
DISPLAY_MATH = re.compile(r'^\$\$(.+)\$\$$')
SCRIPTS = re.compile(r'([_^])\{([^{}]*)\}')
BOLD = re.compile(r'\*\*(.+?)\*\*')


def _set_font(run, latin: str, cjk: str, size, bold: bool = False):
    run.font.name = latin
    run.font.size = size
    run.font.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), cjk)


def _script_runs(para, text: str, size, bold: bool):
    """`x_{0}` / `x^{i}` -> 上下标 run。"""
    pos = 0
    for m in SCRIPTS.finditer(text):
        if m.start() > pos:
            r = para.add_run(text[pos:m.start()])
            _set_font(r, BODY_FONT_LATIN, BODY_FONT_CJK, size, bold)
        r = para.add_run(m.group(2))
        _set_font(r, BODY_FONT_LATIN, BODY_FONT_CJK, size, bold)
        if m.group(1) == '_':
            r.font.subscript = True
        else:
            r.font.superscript = True
        pos = m.end()
    if pos < len(text):
        r = para.add_run(text[pos:])
        _set_font(r, BODY_FONT_LATIN, BODY_FONT_CJK, size, bold)


def _inline(para, text: str, size=BODY_PT):
    """`**bold**` 与上下标混排。"""
    pos = 0
    for m in BOLD.finditer(text):
        if m.start() > pos:
            _script_runs(para, text[pos:m.start()], size, False)
        _script_runs(para, m.group(1), size, True)
        pos = m.end()
    if pos < len(text):
        _script_runs(para, text[pos:], size, False)


def add_body(doc, text: str, indent=True):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = LINE_SPACING
    pf.space_after = Pt(3)
    if indent:
        pf.first_line_indent = FIRST_INDENT
    _inline(p, text)
    return p


def add_heading(doc, text: str, level: int):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before, pf.space_after = Pt(13), Pt(8)
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    _set_font(r, HEAD_FONT_LATIN, HEAD_FONT_CJK, HEAD_PT.get(level, Pt(12)), bold=True)
    return p


def add_formula(doc, latex: str):
    """独立公式行 -> 居中的原生 OMML 公式域。"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.line_spacing = LINE_SPACING
    pf.space_before = pf.space_after = Pt(5)
    try:
        p._p.append(latex_to_omml(latex))
    except Exception as exc:                       # 转换失败时退化为纯文本，不中断出稿
        print(f'  WARN 公式转换失败，已按文本输出: {latex[:48]}… ({exc})', file=sys.stderr)
        _inline(p, latex)
    return p


def add_caption(doc, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    _inline(p, text, CAPTION_PT)


def add_table(doc, rows: list[list[str]]):
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    t.style = 'Table Grid'
    for ri, cells in enumerate(rows):
        for ci, text in enumerate(cells):
            cell = t.cell(ri, ci)
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if ci == 0 else WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(0)
            _inline(p, text.replace('**', ''), TABLE_PT)
            if ri == 0:
                shd = cell._tc.get_or_add_tcPr()
                el = shd.makeelement(qn('w:shd'), {qn('w:val'): 'clear', qn('w:fill'): HEADER_FILL})
                shd.append(el)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return t


def set_orientation(section, landscape: bool):
    if landscape:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = PAGE_H, PAGE_W
    else:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width, section.page_height = PAGE_W, PAGE_H
    for attr in ('top_margin', 'bottom_margin', 'left_margin', 'right_margin'):
        setattr(section, attr, MARGIN)


def add_image(doc, path: Path, landscape: bool, alt: str = ''):
    from PIL import Image
    with Image.open(path) as im:
        w, h = im.size
    max_w, max_h = FIT_MM['landscape' if landscape else 'portrait']
    k = min(max_w / (w / 96 * 25.4), max_h / (h / 96 * 25.4), 1.0)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before, p.paragraph_format.space_after = Pt(10), Pt(4)
    shape = p.add_run().add_picture(str(path), width=Mm(w / 96 * 25.4 * k))
    if alt:
        # Markdown 的 ![alt](path) 里的 alt，写入 docPr 的 name/descr
        doc_pr = shape._inline.docPr
        doc_pr.set('name', alt)
        doc_pr.set('descr', alt)


def render(src: Path, out: Path, landscape_stems: set[str]) -> None:
    doc = Document()
    st = doc.styles['Normal']
    st.font.name, st.font.size = BODY_FONT_LATIN, BODY_PT
    st.element.rPr.rFonts.set(qn('w:eastAsia'), BODY_FONT_CJK)
    set_orientation(doc.sections[0], False)
    cur_landscape = False

    lines = src.read_text(encoding='utf8').split('\n')
    i = 0
    while i < len(lines):
        t = lines[i].strip()
        if not t:
            i += 1
            continue

        if t == '---':
            doc.add_page_break()
            i += 1
            continue

        if t.startswith('|'):                       # Markdown 表格块
            block = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                block.append(lines[i].strip())
                i += 1
            rows = [r.strip('|').split('|') for r in block
                    if not re.fullmatch(r'\|[\s:|-]+\|', r)]
            add_table(doc, [[c.strip() for c in r] for r in rows])
            continue

        m = DISPLAY_MATH.match(t)
        if m:
            add_formula(doc, m.group(1).strip())
            i += 1
            continue

        m = MD_IMG.match(t)
        if m:
            alt, path = m.group(1), (src.parent / m.group(2)).resolve()
            want = path.stem in landscape_stems
            if want != cur_landscape:
                set_orientation(doc.add_section(), want)
                cur_landscape = want
            add_image(doc, path, want, alt)
            i += 1
            continue

        if MD_CAPTION.match(t):
            add_caption(doc, t)
            i += 1
            continue

        m = MD_HEAD.match(t)
        if m:
            add_heading(doc, m.group(2), len(m.group(1)))
            i += 1
            continue

        add_body(doc, t)
        i += 1

    doc.save(out)


def main() -> int:
    ap = argparse.ArgumentParser(description='交底书 Markdown → docx（公式为原生 OMML）')
    ap.add_argument('source')
    ap.add_argument('-o', '--output', required=True)
    ap.add_argument('--landscape', nargs='*', default=['fig1', 'fig2'],
                    help='需要横向页的图片文件名主干，默认 fig1 fig2')
    a = ap.parse_args()

    src, out = Path(a.source), Path(a.output)
    render(src, out, set(a.landscape))

    import zipfile
    xml = zipfile.ZipFile(out).read('word/document.xml').decode('utf8')
    n_img = xml.count('<w:drawing>')
    print(f'已写出 {out} （{out.stat().st_size:,} 字节）')
    print(f'  原生公式域 {xml.count("<m:oMath")} 个 / 内嵌图 {n_img} 张 / '
          f'表格 {xml.count("<w:tbl>")} 个')

    # 有附图说明却一张图都没有，多半是漏了阶段 6，不是有意为之
    md = src.read_text(encoding='utf8')
    if n_img == 0 and re.search(r'(?m)^#*\s*附图说明\s*$', md):
        print('  WARN 文中有「附图说明」但未内嵌任何图片——附图是否漏画？', file=sys.stderr)
    return 0


if __name__ == '__main__':
    sys.exit(main())
